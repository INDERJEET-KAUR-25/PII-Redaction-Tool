from docx import Document


class DocumentProcessor:
    """
    Handles reading, modifying and saving DOCX files.
    """

    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)

    def get_paragraphs(self):
        """
        Returns paragraph objects (not just text).
        """
        return self.document.paragraphs

    def get_tables(self):
        """
        Returns table objects.
        """
        return self.document.tables

    def save(self, output_path):
        self.document.save(output_path)